"""
extractors.py — Extract Sanskrit text from various file formats.

Supports:
  - PDF (text-based — direct extraction; scanned — OCR via pdf2image)
  - DOCX (Word documents)
  - XLSX / XLS (Excel spreadsheets; user picks sloka column)
  - TXT (plain text)
  - Image files (handled via ocr.py)
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

import pandas as pd
from PIL import Image


def extract_from_txt(file_bytes: bytes) -> str:
    """Plain text file — just decode as UTF-8."""
    try:
        return file_bytes.decode('utf-8').strip()
    except UnicodeDecodeError:
        return file_bytes.decode('utf-8', errors='replace').strip()


def extract_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a .docx file."""
    from docx import Document
    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return '\n'.join(paragraphs).strip()


def extract_from_pdf(file_bytes: bytes, use_ocr_if_needed: bool = True) -> str:
    """
    Extract text from a PDF.
      1. Try PyPDF2 for text-based PDFs (fast, lossless)
      2. If the result contains little Devanagari, fall back to OCR on rasterized pages
    """
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(file_bytes))
    direct_text = []
    for page in reader.pages:
        try:
            direct_text.append(page.extract_text() or '')
        except Exception:
            direct_text.append('')
    text = '\n'.join(direct_text).strip()

    # Check whether the direct extraction produced enough Devanagari
    dev_chars = sum(1 for c in text if '\u0900' <= c <= '\u097F')

    if dev_chars < 20 and use_ocr_if_needed:
        # Probably a scanned PDF — fall back to OCR on rasterized pages
        try:
            from pdf2image import convert_from_bytes
            from utils.ocr import extract_text_from_image

            images = convert_from_bytes(file_bytes, dpi=300)
            ocr_texts = []
            for img in images:
                ocr_texts.append(extract_text_from_image(img))
            return '\n'.join(ocr_texts).strip()
        except Exception as e:
            # pdf2image may not be available; return whatever we got
            return text

    return text


def extract_from_excel(
    file_bytes: bytes,
    sheet_name: str | int = 0,
    column: str | None = None,
) -> tuple[list[str], list[str]]:
    """
    Read an Excel file. Returns (extracted_lines, list_of_columns).

    If `column` is None, returns the list of columns for the user to choose.
    If a column is provided, returns the Sanskrit entries from that column.
    """
    df = pd.read_excel(BytesIO(file_bytes), sheet_name=sheet_name)
    columns = list(df.columns)

    if column is None or column not in df.columns:
        return [], columns

    # Pull non-empty string entries from the chosen column
    lines = []
    for val in df[column].dropna():
        txt = str(val).strip()
        if txt:
            lines.append(txt)
    return lines, columns


def extract_from_image(file_bytes: bytes) -> str:
    """Open an image and run OCR."""
    from utils.ocr import extract_text_from_image
    img = Image.open(BytesIO(file_bytes))
    return extract_text_from_image(img)


# ---------------------------------------------------------------------------
# Split extracted text into individual ślokas
# ---------------------------------------------------------------------------
def split_into_slokas(text: str) -> list[str]:
    """
    Split a block of Sanskrit text into individual ślokas.

    Heuristic:
      - Split on double daṇḍa (॥) or single daṇḍa (।) followed by number
      - Split on line breaks if present
      - Filter out fragments too short to be real content
    """
    if not text:
        return []

    # Normalize whitespace and line breaks
    text = re.sub(r'\r\n?', '\n', text)

    # Primary split: on double-danda with optional number (॥1॥, ॥२॥, etc.)
    # or on single danda followed by newline
    pieces = re.split(r'॥[\d०-९\s]*॥|।\s*\n+', text)

    # If that didn't split (no dandas), fall back to line-based splitting
    if len(pieces) == 1:
        pieces = [p for p in text.split('\n') if p.strip()]

    # Clean and filter
    result = []
    for p in pieces:
        # Strip trailing danda / whitespace / Arabic digits / brackets
        p = re.sub(r'[।॥\s]*$', '', p).strip()
        p = re.sub(r'^\s*[\(\[]?\s*\d+\s*[\)\].]?\s*', '', p)  # leading "1." or "(1)"
        # Count Devanagari chars to ensure it's a real śloka fragment
        dev_count = sum(1 for c in p if '\u0900' <= c <= '\u097F')
        if dev_count >= 8:
            result.append(p)

    return result


# ---------------------------------------------------------------------------
# Top-level dispatcher
# ---------------------------------------------------------------------------
def extract_text(
    file_bytes: bytes,
    filename: str,
    excel_sheet: int | str = 0,
    excel_column: str | None = None,
) -> tuple[str, dict]:
    """
    Dispatch to the right extractor based on filename extension.

    Returns (extracted_text, metadata_dict)
    metadata_dict may contain 'columns' (for Excel) so the UI can prompt for
    column selection before running extraction for real.
    """
    ext = Path(filename).suffix.lower().strip('.')
    metadata: dict = {'extension': ext}

    if ext == 'txt':
        return extract_from_txt(file_bytes), metadata

    if ext == 'docx':
        return extract_from_docx(file_bytes), metadata

    if ext == 'pdf':
        return extract_from_pdf(file_bytes), metadata

    if ext in ('xlsx', 'xls'):
        lines, cols = extract_from_excel(file_bytes, excel_sheet, excel_column)
        metadata['columns'] = cols
        if not lines:
            return '', metadata
        return '\n'.join(lines), metadata

    if ext in ('png', 'jpg', 'jpeg', 'bmp', 'tif', 'tiff', 'webp'):
        return extract_from_image(file_bytes), metadata

    raise ValueError(f"Unsupported file type: .{ext}")
